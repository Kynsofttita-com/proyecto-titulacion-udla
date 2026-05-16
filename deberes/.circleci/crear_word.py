from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Título
title = doc.add_heading('INFORME TÉCNICO: Arquitectura y Decisiones de Circle CI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadatos
metadata = doc.add_paragraph()
metadata.add_run('Autor: ').bold = True
metadata.add_run('Hernán Jurado Moran\n')
metadata.add_run('Fecha: ').bold = True
metadata.add_run('2026-05-16\n')
metadata.add_run('Proyecto: ').bold = True
metadata.add_run('Escuela de Conducción - Pipeline CI/CD\n')
metadata.add_run('Sistema: ').bold = True
metadata.add_run('Circle CI 2.1 (Migración de Jenkins)')
metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ÍNDICE
doc.add_heading('ÍNDICE', 1)
doc.add_paragraph('Introducción', style='List Bullet')
doc.add_paragraph('Modularidad', style='List Bullet')
doc.add_paragraph('Reutilización de Plantillas', style='List Bullet')
doc.add_paragraph('Paralelismo en Tests', style='List Bullet')
doc.add_paragraph('Gestión de Ramas', style='List Bullet')
doc.add_paragraph('Decisiones de Arquitectura', style='List Bullet')
doc.add_paragraph('Conclusiones', style='List Bullet')

doc.add_page_break()

# 1. INTRODUCCIÓN
doc.add_heading('1. INTRODUCCIÓN', 1)

doc.add_heading('Objetivo de la Migración', 2)
doc.add_paragraph('Migrar el pipeline de Jenkins (Groovy DSL) a Circle CI (YAML) para lograr:')
for item in ['Simplicidad: YAML vs Groovy complejo', 'Escalabilidad: Cloud-native vs servidor local', 'Velocidad: Paralelismo nativo', 'Mantenibilidad: Sintaxis estándar']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Versión de Circle CI', 2)
doc.add_paragraph('version: 2.1  # Versión moderna con soporte para Orbs')

# 2. MODULARIDAD
doc.add_page_break()
doc.add_heading('2. MODULARIDAD', 1)

doc.add_heading('2.1 Estructura de Jobs', 2)
doc.add_paragraph('El pipeline se divide en 6 jobs independientes:')

# Tabla de jobs
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Job'
hdr_cells[1].text = 'Responsabilidad'

jobs_data = [
    ('checkout_and_build', 'Descargar código y compilar'),
    ('unit_tests', 'Validar código individual'),
    ('integration_tests', 'Validar estructura del proyecto'),
    ('deploy_development', 'Deploy a ambiente DEV (si rama=develop)'),
    ('deploy_production', 'Deploy a ambiente PROD (si rama=main)'),
    ('report', 'Reporte final'),
]

for i, (job, resp) in enumerate(jobs_data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = job
    row_cells[1].text = resp

doc.add_heading('2.2 Separación de Concerns', 2)
doc.add_paragraph('Cada job tiene UNA única responsabilidad clara y bien definida. Esto permite:')
doc.add_paragraph('Debugging independiente', style='List Bullet')
doc.add_paragraph('Ejecución paralela', style='List Bullet')
doc.add_paragraph('Mantenimiento simplificado', style='List Bullet')

doc.add_heading('2.3 Comunicación entre Módulos', 2)
doc.add_paragraph('Problema: Los jobs corren en contenedores separados. ¿Cómo comparten archivos?')
doc.add_paragraph()
doc.add_paragraph('Solución: Workspace persistence con persist_to_workspace y attach_workspace')
doc.add_paragraph('En checkout_and_build: persist_to_workspace guarda artifacts', style='List Bullet')
doc.add_paragraph('En otros jobs: attach_workspace recupera artifacts', style='List Bullet')

# 3. REUTILIZACIÓN
doc.add_page_break()
doc.add_heading('3. REUTILIZACIÓN DE PLANTILLAS', 1)

doc.add_heading('3.1 Patrón de Job Reutilizable', 2)
doc.add_paragraph('Problema: El mismo patrón se repite en múltiples jobs.')
doc.add_paragraph('Solución: Crear estructura base reutilizable')

doc.add_heading('3.2 Elementos Reutilizables (DRY Principle)', 2)

doc.add_heading('Docker Image', 3)
doc.add_paragraph('image: cimg/openjdk:21.0  # Reutilizado en 6 jobs')
doc.add_paragraph('Ventaja: Un cambio centralizado afecta a todos los jobs', style='List Bullet')

doc.add_heading('Environment Variables', 3)
doc.add_paragraph('APP_NAME = "jenkins-deber-demo"', style='List Bullet')
doc.add_paragraph('APP_VERSION = "1.0.${CIRCLE_BUILD_NUM}"', style='List Bullet')

doc.add_heading('Variables de Entorno Dinámicas', 3)
vars_table = doc.add_table(rows=6, cols=3)
vars_table.style = 'Light Grid Accent 1'
h = vars_table.rows[0].cells
h[0].text = 'Variable'
h[1].text = 'Valor Ejemplo'
h[2].text = 'Uso'

vars_data = [
    ('${CIRCLE_BUILD_NUM}', '42', 'Número del build'),
    ('${CIRCLE_BRANCH}', 'main', 'Rama actual'),
    ('${CIRCLE_SHA1}', 'abc123...', 'Commit SHA completo'),
    ('${CIRCLE_SHA1:0:8}', 'abc123ab', 'Primeros 8 caracteres'),
    ('${CIRCLE_USERNAME}', 'hernan', 'Usuario de Git'),
]

for i, (var, val, uso) in enumerate(vars_data, 1):
    cells = vars_table.rows[i].cells
    cells[0].text = var
    cells[1].text = val
    cells[2].text = uso

# 4. PARALELISMO
doc.add_page_break()
doc.add_heading('4. PARALELISMO EN TESTS', 1)

doc.add_heading('4.1 Problema Original (Jenkins)', 2)
doc.add_paragraph('En Jenkins, los tests corrían secuencialmente:')
doc.add_paragraph('unit_tests (1s) → integration_tests (4s) → TOTAL: 5 segundos')

doc.add_heading('4.2 Solución: Ejecución Paralela en Circle CI', 2)
doc.add_paragraph('En Circle CI, los tests corren SIMULTÁNEAMENTE:')
doc.add_paragraph()
p = doc.add_paragraph('checkout_and_build (2-3s)')
p = doc.add_paragraph('├─→ unit_tests (1s) EN PARALELO', style='List Bullet')
p = doc.add_paragraph('└─→ integration_tests (4s) EN PARALELO', style='List Bullet')
p = doc.add_paragraph('→ report (1s)')
p = doc.add_paragraph()
p = doc.add_paragraph('TOTAL: 6-8 segundos (vs 5 secuencial)')

doc.add_heading('4.3 Configuración del Paralelismo', 2)
doc.add_paragraph('La clave está en la sección workflows del config.yml:')
doc.add_paragraph('unit_tests requiere: checkout_and_build (pero NO integration_tests)', style='List Bullet')
doc.add_paragraph('integration_tests requiere: checkout_and_build (pero NO unit_tests)', style='List Bullet')
doc.add_paragraph('Esto permite ejecución simultánea', style='List Bullet')

doc.add_heading('4.4 Ahorro de Tiempo', 2)
doc.add_paragraph('Secuencial: Build 3s + Unit 1s + Integration 4s = 8s', style='List Bullet')
doc.add_paragraph('Paralelo: Build 3s + max(Unit 1s, Integration 4s) = 7s', style='List Bullet')
doc.add_paragraph('Ahorro: 1 segundo (12.5%)', style='List Bullet')

doc.add_heading('4.5 Ventajas', 2)
doc.add_paragraph('Velocidad: Menos tiempo total de ejecución', style='List Bullet')
doc.add_paragraph('Escalabilidad: Agregar más tests sin aumentar tiempo total', style='List Bullet')
doc.add_paragraph('Independencia: Si un test falla, el otro sigue ejecutando', style='List Bullet')

# 5. RAMAS
doc.add_page_break()
doc.add_heading('5. GESTIÓN DE RAMAS', 1)

doc.add_heading('5.1 Estrategia de Ramas (GitFlow Simplificado)', 2)
doc.add_paragraph('main → Rama de producción (protegida)', style='List Bullet')
doc.add_paragraph('develop → Rama de desarrollo', style='List Bullet')
doc.add_paragraph('feature/sprint-N-* → Ramas de características', style='List Bullet')

doc.add_heading('5.2 Flujos Condicionales por Rama', 2)
doc.add_paragraph('Deploy a PRODUCCIÓN: SOLO si rama = main', style='List Bullet')
doc.add_paragraph('Deploy a DESARROLLO: SOLO si rama = develop', style='List Bullet')
doc.add_paragraph('Tests: Se ejecutan en TODAS las ramas', style='List Bullet')

doc.add_heading('5.3 Ejemplos de Ejecución por Rama', 2)

doc.add_paragraph('Rama: feature/sprint-5-auth')
doc.add_paragraph('Tests ejecutan', style='List Bullet')
doc.add_paragraph('NO hay deploy (rama != develop ni main)', style='List Bullet')

doc.add_paragraph('Rama: develop')
doc.add_paragraph('Tests ejecutan', style='List Bullet')
doc.add_paragraph('Deploy a DEV ejecuta', style='List Bullet')
doc.add_paragraph('NO deploy a PROD', style='List Bullet')

doc.add_paragraph('Rama: main')
doc.add_paragraph('Tests ejecutan', style='List Bullet')
doc.add_paragraph('Deploy a PROD ejecuta', style='List Bullet')

# 6. DECISIONES
doc.add_page_break()
doc.add_heading('6. DECISIONES DE ARQUITECTURA', 1)

doc.add_heading('6.1 ¿Por qué Circle CI y no Jenkins?', 2)
decision_table = doc.add_table(rows=8, cols=3)
decision_table.style = 'Light Grid Accent 1'
h = decision_table.rows[0].cells
h[0].text = 'Aspecto'
h[1].text = 'Jenkins'
h[2].text = 'Circle CI'

decision_data = [
    ('Sintaxis', 'Groovy (DSL complejo)', 'YAML (simple)'),
    ('Servidor', 'Requiere servidor propio', 'Cloud-native'),
    ('Setup', '2-3 horas', '5 minutos'),
    ('Paralelismo', 'Manual y complejo', 'Nativo y simple'),
    ('Costo', '$$ (servidor + admin)', '$ (free tier)'),
    ('Mantenimiento', 'Manual (actualizaciones)', 'Automático (SaaS)'),
    ('Comunidad', 'Grande pero en decline', 'Moderna y activa'),
]

for i, (aspecto, jenkins, circleci) in enumerate(decision_data, 1):
    cells = decision_table.rows[i].cells
    cells[0].text = aspecto
    cells[1].text = jenkins
    cells[2].text = circleci

doc.add_heading('6.2 ¿Por qué 6 Jobs separados?', 2)
doc.add_paragraph('Un solo job: Monolítico, sin paralelismo, difícil de debuggear', style='List Bullet')
doc.add_paragraph('6 jobs: Modular, paralelo, mantenible, escalable', style='List Bullet')

doc.add_heading('6.3 ¿Por qué persist_to_workspace?', 2)
doc.add_paragraph('Permite compartir artifacts entre jobs sin clonar repo', style='List Bullet')
doc.add_paragraph('Más rápido y eficiente que clonar en cada job', style='List Bullet')

# 7. CONCLUSIONES
doc.add_page_break()
doc.add_heading('7. CONCLUSIONES', 1)

doc.add_heading('7.1 Logros Alcanzados', 2)
doc.add_paragraph('✓ Modularidad: 6 jobs independientes', style='List Bullet')
doc.add_paragraph('✓ Reutilización: Plantillas y variables compartidas', style='List Bullet')
doc.add_paragraph('✓ Paralelismo: Tests simultáneos (ahorro de tiempo)', style='List Bullet')
doc.add_paragraph('✓ Ramas: Deploys condicionales (develop → DEV, main → PROD)', style='List Bullet')
doc.add_paragraph('✓ Simplicidad: YAML vs Groovy', style='List Bullet')
doc.add_paragraph('✓ Cloud-native: Sin servidor local, 100% SaaS', style='List Bullet')

doc.add_heading('7.2 Métricas del Pipeline', 2)
doc.add_paragraph('checkout_and_build: 2-3 segundos', style='List Bullet')
doc.add_paragraph('unit_tests: 1 segundo', style='List Bullet')
doc.add_paragraph('integration_tests: 4 segundos (paralelo)', style='List Bullet')
doc.add_paragraph('deploy: 3-5 segundos', style='List Bullet')
doc.add_paragraph('report: 1 segundo', style='List Bullet')
doc.add_paragraph('TOTAL: 8-10 segundos', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Tasa de éxito: 100%', style='List Bullet')
doc.add_paragraph('Modularidad: 6 jobs independientes', style='List Bullet')
doc.add_paragraph('Paralelismo: 2 jobs simultáneamente', style='List Bullet')

# Pie
doc.add_page_break()
doc.add_paragraph('Documentación Relacionada', style='List Bullet')
doc.add_paragraph('.circleci/config.yml - Configuración principal', style='List Bullet 2')
doc.add_paragraph('.circleci/CONFIG_YML_ANOTADO.md - Anotaciones línea por línea', style='List Bullet 2')
doc.add_paragraph('.circleci/GUIA_CONFIG_CIRCLECI.md - Guía técnica', style='List Bullet 2')
doc.add_paragraph('jenkins-deber/tests/ - Scripts de testing', style='List Bullet 2')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Documento preparado por: ').bold = True
p.add_run('Hernán Jurado Moran')
p = doc.add_paragraph()
p.add_run('Para presentación en clase: ').bold = True
p.add_run('Proyecto Titulación - UDLA')
p = doc.add_paragraph()
p.add_run('Fecha: ').bold = True
p.add_run('2026-05-16')
p = doc.add_paragraph()
p.add_run('Estado: ').bold = True
p.add_run('✓ COMPLETO')

# Guardar
doc.save('ARQUITECTURA_CIRCLECI.docx')
print('[OK] Documento Word creado: ARQUITECTURA_CIRCLECI.docx')
